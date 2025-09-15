"""
Google Document AI Text Extraction Module
Replaces PyPDF2/pytesseract OCR functionality with Google Document AI
"""

import io
import logging
from typing import Optional, Dict, Any
from pathlib import Path

try:
    from google.cloud import documentai_v1 as documentai
    from google.oauth2 import service_account
except ImportError:
    documentai = None
    service_account = None

import google.generativeai as genai
from .config import Config


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentAIExtractor:
    """Extract text from documents using Google Document AI"""
    
    def __init__(self):
        self.config = Config()
        self.client = None
        self._initialize_client()
    
    def _initialize_client(self):
        """Initialize Document AI client"""
        try:
            if documentai is None:
                logger.warning("Google Cloud Document AI not available. Using fallback methods.")
                return
                
            # Set up client options for the specified location
            client_options = {
                "api_endpoint": f"{self.config.DOCUMENT_AI_LOCATION}-documentai.googleapis.com"
            }
            self.client = documentai.DocumentProcessorServiceClient(
                client_options=client_options
            )
            logger.info("Document AI client initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize Document AI client: {e}")
            self.client = None
    
    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF using Document AI OCR
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            if self.client is None or not self.config.DOCUMENT_AI_PROCESSOR_ID:
                return self._fallback_pdf_extraction(file_path)
            
            # Read the file
            with open(file_path, "rb") as document_file:
                document_content = document_file.read()
            
            # Configure the process request
            raw_document = documentai.RawDocument(
                content=document_content, 
                mime_type="application/pdf"
            )
            
            # The full resource name of the processor
            processor_name = self.client.processor_path(
                self.config.GOOGLE_CLOUD_PROJECT_ID,
                self.config.DOCUMENT_AI_LOCATION,
                self.config.DOCUMENT_AI_PROCESSOR_ID
            )
            
            request = documentai.ProcessRequest(
                name=processor_name,
                raw_document=raw_document
            )
            
            # Process the document
            result = self.client.process_document(request=request)
            document = result.document
            
            # Extract text and metadata
            extracted_data = {
                'text': document.text,
                'pages': len(document.pages),
                'entities': self._extract_entities_from_document(document),
                'tables': self._extract_tables_from_document(document),
                'confidence': self._calculate_confidence(document)
            }
            
            logger.info(f"Successfully extracted text from {file_path} using Document AI")
            return extracted_data
            
        except Exception as e:
            logger.error(f"Document AI extraction failed: {e}")
            return self._fallback_pdf_extraction(file_path)
    
    def _fallback_pdf_extraction(self, file_path: str) -> Dict[str, Any]:
        """Fallback PDF extraction using PyPDF2"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    'text': text.strip(),
                    'pages': len(pdf_reader.pages),
                    'entities': [],
                    'tables': [],
                    'confidence': 0.8  # Estimated confidence for fallback
                }
                
        except Exception as e:
            logger.error(f"Fallback PDF extraction failed: {e}")
            return {'text': '', 'pages': 0, 'entities': [], 'tables': [], 'confidence': 0.0}
    
    def extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'text': text.strip(),
                'pages': 1,  # DOCX doesn't have traditional pages
                'entities': [],
                'tables': tables,
                'confidence': 0.95
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return {'text': '', 'pages': 0, 'entities': [], 'tables': [], 'confidence': 0.0}
    
    def _extract_entities_from_document(self, document) -> list:
        """Extract entities from Document AI response"""
        entities = []
        
        for entity in document.entities:
            entities.append({
                'type': entity.type_,
                'text': entity.text_anchor.content if entity.text_anchor else '',
                'confidence': entity.confidence
            })
        
        return entities
    
    def _extract_tables_from_document(self, document) -> list:
        """Extract tables from Document AI response"""
        tables = []
        
        for page in document.pages:
            for table in page.tables:
                table_data = []
                for row in table.header_rows + table.body_rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = ""
                        for segment in cell.layout.text_anchor.text_segments:
                            start_index = segment.start_index
                            end_index = segment.end_index
                            cell_text += document.text[start_index:end_index]
                        row_data.append(cell_text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
        
        return tables
    
    def _calculate_confidence(self, document) -> float:
        """Calculate overall confidence score"""
        if not document.pages:
            return 0.0
        
        total_confidence = 0.0
        total_elements = 0
        
        for page in document.pages:
            for block in page.blocks:
                if hasattr(block, 'confidence'):
                    total_confidence += block.confidence
                    total_elements += 1
        
        return total_confidence / total_elements if total_elements > 0 else 0.8

class GeminiTextExtractor:
    """Alternative text extraction using Gemini API for documents"""
    
    def __init__(self):
        self.config = Config()
        genai.configure(api_key=self.config.GEMINI_API_KEY)
        self.model = genai.GenerativeModel(self.config.GEMINI_FLASH_MODEL)
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text using Gemini API file processing
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Upload file to Gemini
            uploaded_file = genai.upload_file(path=file_path)
            
            # Wait for file processing
            while uploaded_file.state.name == "PROCESSING":
                time.sleep(1)
                uploaded_file = genai.get_file(uploaded_file.name)
            
            if uploaded_file.state.name == "FAILED":
                raise Exception("File processing failed")
            
            # Extract text using Gemini
            prompt = "Extract all text content from this document. Preserve formatting and structure."
            response = self.model.generate_content([uploaded_file, prompt])
            
            return {
                'text': response.text,
                'pages': 1,  # Gemini doesn't provide page count
                'entities': [],
                'tables': [],
                'confidence': 0.9
            }
            
        except Exception as e:
            logger.error(f"Gemini text extraction failed: {e}")
            return {'text': '', 'pages': 0, 'entities': [], 'tables': [], 'confidence': 0.0}
        finally:
            # Clean up uploaded file
            try:
                genai.delete_file(uploaded_file.name)
            except:
                pass